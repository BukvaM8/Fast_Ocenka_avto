# -----------------------------
# АВТОУСТАНОВКА ЗАВИСИМОСТЕЙ ИЗ requirements.txt (жёстко, без --user)
# -----------------------------
import subprocess, sys, os, tempfile
import streamlit as st

def install_requirements_strict(req_path: str):
    env = os.environ.copy()
    # игнорируем пользовательские pip-конфиги
    empty_cfg = os.path.join(tempfile.gettempdir(), "empty_pip.cfg")
    if not os.path.exists(empty_cfg):
        with open(empty_cfg, "w", encoding="utf-8") as f:
            f.write("")
    env["PIP_CONFIG_FILE"] = empty_cfg
    env["PIP_DISABLE_PIP_VERSION_CHECK"] = "1"
    env.pop("PIP_USER", None)
    env.pop("PIP_TARGET", None)
    env.pop("PYTHONUSERBASE", None)

    cmd = [sys.executable, "-m", "pip", "install", "-r", req_path]
    subprocess.check_call(cmd, env=env)

if "reqs_installed" not in st.session_state:
    req_path = os.path.join(os.path.dirname(__file__), "requirements.txt")
    if os.path.exists(req_path):
        try:
            install_requirements_strict(req_path)
            st.session_state["reqs_installed"] = True
        except subprocess.CalledProcessError as e:
            print("Ошибка при установке зависимостей:", e)
            st.error("Не удалось установить зависимости. См. терминал.")
    else:
        st.session_state["reqs_installed"] = True

# -----------------------------
# ИМПОРТЫ И НАСТРОЙКИ
# -----------------------------
import io
import random
import math
from datetime import date
from pathlib import Path

from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

st.set_page_config(page_title="Оценка авто — MVP", layout="centered")


def build_photos_subdoc(doc, files, per_row=2, empty_message="Фотографии не загружены."):
    subdoc = doc.new_subdoc()
    if not files:
        if empty_message:
            subdoc.add_paragraph(empty_message)
        return subdoc
    rows = math.ceil(len(files) / per_row)
    table = subdoc.add_table(rows=rows, cols=per_row)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    idx = 0
    for r in range(rows):
        for c in range(per_row):
            cell = table.cell(r, c)
            for paragraph in cell.paragraphs:
                paragraph.text = ""
            if idx < len(files):
                image_stream = io.BytesIO(files[idx]['data'])
                run = cell.paragraphs[0].add_run()
                run.add_picture(image_stream, width=Inches(3.0))
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                idx += 1
    return subdoc

DEFAULT_CONTRACTOR = "ООО «Агентство «Бизнес-Актив»"
TEMPLATE_NAME = "mers_ocenka.docx"   # имя файла шаблона
TEMPLATES_DIR = Path(__file__).parent / "templates"
GENERATED_DIR = Path(__file__).parent / "generated"
GENERATED_DIR.mkdir(exist_ok=True)

# -----------------------------
# УТИЛИТЫ
# -----------------------------
def generate_uuid7() -> str:
    return f"{random.randint(0, 9_999_999):07d}"

def safe_str_date(d: date) -> str:
    return d.strftime("%d.%m.%Y") if isinstance(d, date) else str(d)

def summarize_attachments(files, failures):
    lines = []
    if files:
        for idx, item in enumerate(files, start=1):
            lines.append(f"{idx}. {item['name']}")
    if failures:
        lines.append("Не удалось загрузить: " + ', '.join(failures))
    if not lines:
        return "Файлы не загружены."
    return '\n'.join(lines)


IMAGE_SUFFIXES = {'.png', '.jpg', '.jpeg', '.bmp', '.gif', '.tif', '.tiff', '.webp'}


def _is_image_file(name: str) -> bool:
    return Path(name or '').suffix.lower() in IMAGE_SUFFIXES


def build_appendix_entries(doc, files, failures):
    from docxtpl import InlineImage
    from docx.shared import Mm

    entries = []

    for item in files:
        display_name = item.get('name', 'без названия')
        if not _is_image_file(display_name):
            continue
        try:
            image_stream = io.BytesIO(item['data'])
            image_stream.seek(0)
            image_stream.name = display_name
            img = InlineImage(doc, image_stream, width=Mm(140))
        except Exception as exc:
            failures.append(f"{display_name} (ошибка вставки: {exc})")
            continue
        entries.append({'image': img})

    return entries


ANALOG_EMPTY_TITLE = "без названия"
MAX_ANALOGS = 10
ANALOG_TITLE_TEMPLATE = "Предложение по продаже транспортного средства (объект-аналог №{index})"


def ensure_analog_state():
    slots = st.session_state.setdefault("analog_slots", [1])
    if "analog_counter" not in st.session_state:
        st.session_state["analog_counter"] = max(slots) if slots else 0


def add_analog_slot():
    ensure_analog_state()
    slots = list(st.session_state["analog_slots"])
    if len(slots) >= MAX_ANALOGS:
        return
    counter = st.session_state.get("analog_counter", 0) + 1
    st.session_state["analog_counter"] = counter
    slots.append(counter)
    st.session_state["analog_slots"] = slots


def remove_analog_slot(slot_id: int):
    slots = [s for s in st.session_state.get("analog_slots", []) if s != slot_id]
    st.session_state["analog_slots"] = slots
    st.session_state.pop(f"analog_title_{slot_id}", None)
    st.session_state.pop(f"analog_files_{slot_id}", None)
    st.session_state.pop(f"analog_source_{slot_id}", None)


def default_analog_heading(index: int) -> str:
    return ANALOG_TITLE_TEMPLATE.format(index=index)


def format_analog_source_text(text: str) -> str:
    cleaned = (text or "").strip()
    if not cleaned:
        return ""
    lower_cleaned = cleaned.lower()
    prefix = "источник информации"
    if lower_cleaned.startswith(prefix):
        return cleaned
    if not cleaned.startswith("(") and not cleaned.endswith(")"):
        cleaned = f"({cleaned})"
    return f"Источник информации: {cleaned}"


def format_analog_source(text: str):
    formatted = format_analog_source_text(text)
    if not formatted:
        return ""
    from docxtpl import RichText
    rt = RichText()
    rt.add(formatted, italic=True)
    return rt


def paragraph_has_drawings(paragraph):
    for run in paragraph.runs:
        if run._element.xpath('.//w:drawing'):
            return True
    return False


def paragraph_has_page_break(paragraph):
    for run in paragraph.runs:
        for br in run._element.xpath('.//w:br'):
            br_type = br.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type')
            if br_type in (None, 'page'):
                return True
    return False


def paragraph_has_content(paragraph):
    return bool(paragraph.text.strip()) or paragraph_has_drawings(paragraph)


def cleanup_unused_analog_pages(doc_path, analog_count):
    doc = Document(doc_path)
    import re
    paragraphs = doc.paragraphs
    analog_pattern = re.compile(r'объект-аналог\s*№\s*(\d+)', re.IGNORECASE)
    last_idx = -1
    for idx, paragraph in enumerate(paragraphs):
        match = analog_pattern.search(paragraph.text)
        if match:
            try:
                number = int(match.group(1))
            except ValueError:
                continue
            if number <= analog_count:
                last_idx = idx
    start_idx = last_idx + 1 if analog_count else 0
    removed_any = False
    i = start_idx
    while i < len(doc.paragraphs):
        paragraphs = doc.paragraphs
        if i >= len(paragraphs):
            break
        paragraph = paragraphs[i]
        if paragraph_has_content(paragraph):
            i += 1
            continue
        block_start = i
        block_end = i
        has_break = paragraph_has_page_break(paragraph)
        while block_end + 1 < len(paragraphs):
            next_paragraph = paragraphs[block_end + 1]
            if paragraph_has_content(next_paragraph):
                break
            block_end += 1
            if paragraph_has_page_break(next_paragraph):
                has_break = True
        next_index = block_end + 1
        if has_break and block_start >= start_idx:
            for pos in range(block_end, block_start - 1, -1):
                par = doc.paragraphs[pos]
                parent = par._element.getparent()
                parent.remove(par._element)
            removed_any = True
        else:
            i = next_index
            continue
        # after removal, continue from block_start (same index)
        i = block_start
    if removed_any:
        doc.save(doc_path)

# -----------------------------
# МОДАЛЬНОЕ ОКНО АВТОРИЗАЦИИ
# -----------------------------
@st.dialog("Авторизация")
def login_dialog():
    st.write("Введите ваши данные, чтобы продолжить работу с приложением.")
    name = st.text_input("Имя", key="auth_name")
    login_val = st.text_input("Логин", key="auth_login")

    if login_val and not st.session_state.get("uuid7"):
        st.session_state["uuid7"] = generate_uuid7()
    if not login_val:
        st.session_state.pop("uuid7", None)

    st.text_input(
        "Ваш UUID (7 цифр)",
        value=st.session_state.get("uuid7", ""),
        key="uuid7_display",
        disabled=True,
        help="Генерируется автоматически после ввода логина."
    )

    can_submit = bool(name.strip()) and bool(login_val.strip()) and bool(st.session_state.get("uuid7"))
    if st.button("Войти", type="primary", disabled=not can_submit):
        st.session_state["user_name"] = name.strip()
        st.session_state["user_login"] = login_val.strip()
        st.session_state["auth_ok"] = True
        st.rerun()

# Гейт авторизации
if not st.session_state.get("auth_ok"):
    login_dialog()
    st.stop()

# -----------------------------
# САЙДБАР
# -----------------------------
with st.sidebar:
    st.markdown("### Профиль")
    st.markdown(f"**Имя:** {st.session_state.get('user_name', '')}")
    st.markdown(f"**Логин:** {st.session_state.get('user_login', '')}")
    st.markdown(f"**UUID:** {st.session_state.get('uuid7', '')}")
    if st.button("Выйти"):
        for k in ("auth_ok","user_name","user_login","auth_name","auth_login","uuid7","uuid7_display"):
            st.session_state.pop(k, None)
        st.rerun()

# -----------------------------
# ОСНОВНОЙ ИНТЕРФЕЙС
# -----------------------------
st.title("Оценка авто")

if "contractor" not in st.session_state:
    st.session_state["contractor"] = DEFAULT_CONTRACTOR

ensure_analog_state()

with st.form("auto_appraisal_form", clear_on_submit=False):
    col1, col2 = st.columns(2)

    with col1:
        contract_no   = st.text_input("Номер договора", key="contract_no")
        basis         = st.text_input("Основание", key="basis")
        valuation_date= st.date_input("Дата оценки:", value=date.today(), key="valuation_date")
        report_date   = st.date_input("Дата составления Отчета об оценке:", value=date.today(), key="report_date")
        otchet_number = st.text_input("Номер отчета", key="otchet_number")
        object_type   = st.text_input("Тип оцениваемого объекта", key="object_type")
        car_number    = st.text_input("Регистрационный номер автомобиля", key="car_number")

    with col2:
        car_name      = st.text_input("Наим. транспортного средства", key="car_name")
        vin_model     = st.text_input("VIN", key="vin_model")
        customer      = st.text_input("Заказчик:", key="customer")
        contractor    = st.text_input(
            "Исполнитель:",
            key="contractor",
            help="По умолчанию подставляется ООО «Агентство «Бизнес-Актив», измени при необходимости."
        )
        price_no_vat  = st.number_input("Стоимость без НДС:", min_value=0.0, step=0.01, format="%.2f", key="price_no_vat")
        price_vat     = st.number_input("Стоимость с НДС:",  min_value=0.0, step=0.01, format="%.2f", key="price_vat")

    object_photos_raw = st.file_uploader("Фотографии объекта оценки", accept_multiple_files=True, key="object_photos")
    appendix_1_files_raw = st.file_uploader("Приложение 1", accept_multiple_files=True, key="appendix_1")
    appendix_2_files_raw = st.file_uploader("Приложение 2", accept_multiple_files=True, key="appendix_2")
    rights_files_raw = st.file_uploader("Подтверждение права оценщика и исполнителя заниматься оценочной деятельностью", accept_multiple_files=True, key="rights_docs")

    with st.expander("Объекты-аналоги", expanded=False):
        slots = st.session_state.get("analog_slots", [])
        st.caption("Добавляйте до 10 объектов-аналогов. Описания и фотографии попадут в шаблон отчета.")
        if not slots:
            st.info("Нажмите «Добавить объект-аналог», чтобы создать первый блок.")
        for display_index, slot_id in enumerate(slots, start=1):
            cols = st.columns([6, 1])
            title_key = f"analog_title_{slot_id}"
            default_value = st.session_state.get(title_key, default_analog_heading(display_index))
            with cols[0]:
                st.text_input(
                    f"Название / описание для аналога №{display_index}",
                    value=default_value,
                    key=title_key,
                    help="Это значение подставится в {{ object_analogN }}.",
                )
            with cols[1]:
                st.form_submit_button(
                    "Удалить",
                    type="secondary",
                    key=f"remove_analog_{slot_id}",
                    on_click=remove_analog_slot,
                    kwargs={"slot_id": slot_id},
                    help="Удалить этот объект-аналог",
                )
            st.file_uploader(
                f"Фотографии для аналога №{display_index}",
                accept_multiple_files=True,
                key=f"analog_files_{slot_id}",
                help="Файлы будут размещены в {{ object_analogN_photo }}.",
            )
            source_key = f"analog_source_{slot_id}"
            default_source = st.session_state.get(source_key, "")
            st.text_input(
                f"Источник для аналога №{display_index}",
                value=default_source,
                key=source_key,
                help="Текст подставится в {{ a_sourceN }} и будет показан курсивом после фотографий.",
                placeholder="https://...",
            )
            if display_index != len(slots):
                st.divider()
        st.form_submit_button(
            "➕ Добавить объект-аналог",
            key="add_analog_button",
            type="secondary",
            on_click=add_analog_slot,
            disabled=len(slots) >= MAX_ANALOGS,
            help="Добавить ещё один блок загрузки",
        )

    submitted = st.form_submit_button("Сформировать и скачать DOCX", type="primary")



if submitted:
    appendix_1_files = []
    appendix_1_failures = []
    for uploaded_file in (appendix_1_files_raw or []):
        try:
            payload = uploaded_file.getvalue()
            if not payload:
                raise ValueError('empty payload')
        except Exception:
            appendix_1_failures.append(uploaded_file.name or 'без названия')
        else:
            appendix_1_files.append({
                'name': uploaded_file.name or 'без названия',
                'data': payload,
                'size': len(payload),
            })

    appendix_2_files = []
    appendix_2_failures = []
    for uploaded_file in (appendix_2_files_raw or []):
        try:
            payload = uploaded_file.getvalue()
            if not payload:
                raise ValueError('empty payload')
        except Exception:
            appendix_2_failures.append(uploaded_file.name or 'без названия')
        else:
            appendix_2_files.append({
                'name': uploaded_file.name or 'без названия',
                'data': payload,
                'size': len(payload),
            })

    rights_files = []
    rights_failures = []
    for uploaded_file in (rights_files_raw or []):
        try:
            payload = uploaded_file.getvalue()
            if not payload:
                raise ValueError('empty payload')
        except Exception:
            rights_failures.append(uploaded_file.name or 'без названия')
        else:
            rights_files.append({
                'name': uploaded_file.name or 'без названия',
                'data': payload,
                'size': len(payload),
            })

    object_photos = []
    object_photo_failures = []
    for uploaded_file in (object_photos_raw or []):
        try:
            payload = uploaded_file.getvalue()
            if not payload:
                raise ValueError('empty payload')
        except Exception:
            object_photo_failures.append(uploaded_file.name or 'без названия')
        else:
            object_photos.append({
                'name': uploaded_file.name or 'без названия',
                'data': payload,
                'size': len(payload),
            })

    analog_slots_snapshot = st.session_state.get('analog_slots', [])
    analog_results = []
    analog_failures = []
    for display_index, slot_id in enumerate(analog_slots_snapshot, start=1):
        title_key = f"analog_title_{slot_id}"
        raw_title = (st.session_state.get(title_key) or '').strip()
        if not raw_title:
            raw_title = default_analog_heading(display_index)
        source_key = f"analog_source_{slot_id}"
        raw_source = (st.session_state.get(source_key) or '').strip()
        files_key = f"analog_files_{slot_id}"
        uploaded_files = st.session_state.get(files_key) or []
        slot_files = []
        slot_failures = []
        for uploaded_file in uploaded_files:
            try:
                payload = uploaded_file.getvalue()
                if not payload:
                    raise ValueError('empty payload')
            except Exception:
                fail_name = uploaded_file.name or ANALOG_EMPTY_TITLE
                slot_failures.append(fail_name)
                analog_failures.append(f"аналог №{display_index}: {fail_name}")
            else:
                slot_files.append({
                    'name': uploaded_file.name or ANALOG_EMPTY_TITLE,
                    'data': payload,
                    'size': len(payload),
                })
        analog_results.append({
            'index': display_index,
            'slot_id': slot_id,
            'title': raw_title,
            'source': raw_source,
            'files': slot_files,
            'failures': slot_failures,
        })

    appendix_1_names = [item['name'] for item in appendix_1_files]
    appendix_2_names = [item['name'] for item in appendix_2_files]
    rights_names = [item['name'] for item in rights_files]
    object_photo_names = [item['name'] for item in object_photos]
    failed_uploads = (
        appendix_1_failures
        + appendix_2_failures
        + rights_failures
        + object_photo_failures
        + analog_failures
    )

    analog_record_entries = [
        {
            'Номер': item['index'],
            'Название': item['title'],
            'Файлы': [file['name'] for file in item['files']],
            'Ошибки': item['failures'],
            'Источник': format_analog_source_text(item['source']) if (item['files'] and item['source']) else '',
        }
        for item in analog_results
    ]

    # Собираем запись
    record = {
        "user_uuid": st.session_state.get("uuid7"),
        "user_name": st.session_state.get("user_name"),
        "user_login": st.session_state.get("user_login"),
        "Номер договора": contract_no,
        "Основание": basis,
        "Дата оценки": safe_str_date(valuation_date),
        "Дата составления отчета": safe_str_date(report_date),
        "Заказчик": customer,
        "Подрядчик": contractor,
        "Стоимость с НДС": price_vat,
        "Стоимость без НДС": price_no_vat,
        "Номер отчёта": otchet_number,
        "Название ТС": object_type,
        "Регистрационный номер автомобиля": car_number,
        "Доп. наименование ТС": car_name,
        "VIN": vin_model,
        "Приложение 1 (файлы)": appendix_1_names,
        # "Приложение 1 (ошибки)": appendix_1_failures,
        "Приложение 2 (файлы)": appendix_2_names,
        # "Приложение 2 (ошибки)": appendix_2_failures,
        "Подтверждение права (файлы)": rights_names,
        # "Подтверждение права (ошибки)": rights_failures,
        "Фотографии объекта": object_photo_names,
        "Объекты-аналоги": analog_record_entries,
    }
    st.success("Данные сохранены (локально в сессии).")
    with st.expander("Проверить данные перед подстановкой в шаблон"):
        st.json(record)

    # ---- РЕНДЕР ДОКУМЕНТА ИЗ ШАБЛОНА ----
    tpl_path = TEMPLATES_DIR / TEMPLATE_NAME
    if not tpl_path.exists():
        st.error(
            f"Не найден шаблон: {tpl_path}\n"
            f"Создай его и положи в папку templates. "
            f"В следующем шаге добавим файл."
        )
    else:
        # Импортируем тут, чтобы приложение работало даже без docxtpl до клика
        try:
            from docxtpl import DocxTemplate
        except Exception as e:
            st.error(
                "Не удалось импортировать docxtpl. "
                "Убедись, что в requirements.txt есть строка: docxtpl"
            )
            st.stop()

        # МАППИНГ ПОЛЕЙ -> МЕТКИ {{ ... }} В ШАБЛОНЕ
        # (использую именно те ключи, которые ты перечислил)
        doc = DocxTemplate(str(tpl_path))

        appendix_1_entries = build_appendix_entries(doc, appendix_1_files, appendix_1_failures)
        appendix_2_entries = build_appendix_entries(doc, appendix_2_files, appendix_2_failures)
        rights_entries = build_appendix_entries(doc, rights_files, rights_failures)

        context = {
            # как ты указал: Основание -> {{ contract_number }} (да, дублируется)
            "contract_number": contract_no,               # Номер контракта / и по твоей строке "Основание"
            "date_ocenka": safe_str_date(valuation_date),
            "date_otcheta": safe_str_date(report_date),
            "customer_name": customer,

            # нижний блок:
            "contractor": contractor,
            "otchet_number": otchet_number,
            "object_type": object_type,
            "car_name": car_name,
            "vin_model": vin_model,
            "cost_of_assessment": f"{price_no_vat:,.2f}".replace(",", " "),
            "cost_of_assessment_NDS": f"{price_vat:,.2f}".replace(",", " "),
        }
        context["appendix_1_summary"] = summarize_attachments(appendix_1_files, appendix_1_failures)
        context["appendix_2_summary"] = summarize_attachments(appendix_2_files, appendix_2_failures)
        context["rights_summary"] = summarize_attachments(rights_files, rights_failures)
        context["appendix_1_entries"] = appendix_1_entries
        context["appendix_2_entries"] = appendix_2_entries
        context["rights_entries"] = rights_entries
        context["object_ocenki"] = build_photos_subdoc(
            doc,
            object_photos,
            empty_message="Фотографии объекта не загружены."
        )
        for idx in range(1, MAX_ANALOGS + 1):
            if idx <= len(analog_results):
                analog_item = analog_results[idx - 1]
                context[f"object_analog{idx}"] = analog_item['title']
                if analog_item['files'] and analog_item['source']:
                    context[f"a_source{idx}"] = format_analog_source(analog_item['source'])
                else:
                    context[f"a_source{idx}"] = ""
                empty_msg = "Фотографии объекта-аналога не загружены." if not analog_item['files'] else ''
                context[f"object_analog{idx}_photo"] = build_photos_subdoc(
                    doc,
                    analog_item['files'],
                    empty_message=empty_msg
                )
            else:
                context[f"object_analog{idx}"] = ''
                context[f"a_source{idx}"] = ""
                context[f"object_analog{idx}_photo"] = build_photos_subdoc(
                    doc,
                    [],
                    empty_message=''
                )

        doc.render(context)

        out_name = f"Отчёт_{contract_no or 'без_номера'}_{st.session_state.get('uuid7')}.docx"

        out_path = GENERATED_DIR / out_name

        doc.save(str(out_path))

        cleanup_unused_analog_pages(out_path, len(analog_results))




        with open(out_path, "rb") as f:

            data = f.read()



        st.download_button(

            label="⬇️ Скачать сформированный DOCX",

            data=data,

            file_name=out_name,

            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",

            type="primary"

        )

        if failed_uploads:

            st.warning(

                "Не все файлы приложений были загружены: "

                + ', '.join(failed_uploads)

                + ". Проверьте соединение и попробуйте загрузить их повторно."

            )



        st.info(f"Файл также сохранён локально: {out_path}")

        # ----------- БД (оставлено закомментированным) -----------
        # import mysql.connector
        # try:
        #     conn = mysql.connector.connect(
        #         host="localhost",
        #         user="user",
        #         password="password",
        #         database="auto_appraisal_db",
        #     )
        #     cur = conn.cursor()
        #     cur.execute(
        #         """
        #         INSERT INTO appraisals (
        #             user_uuid, user_name, user_login,
        #             contract_no, basis, valuation_date, report_date,
        #             customer, contractor, price_vat, price_no_vat,
        #             otchet_number, object_type, car_name, vin_model, file_name
        #         ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
        #         """,
        #         (
        #             record["user_uuid"], record["user_name"], record["user_login"],
        #             contract_no, basis, safe_str_date(valuation_date), safe_str_date(report_date),
        #             customer, contractor, price_vat, price_no_vat,
        #             otchet_number, object_type, car_name, vin_model, out_name
        #         ),
        #     )
        #     conn.commit()
        #     cur.close()
        #     conn.close()
        #     st.success("Запись сохранена в базу данных MySQL.")
        # except Exception as e:
        #     st.error(f"Ошибка сохранения в MySQL: {e}")
        # ----------------------------------------------------------
